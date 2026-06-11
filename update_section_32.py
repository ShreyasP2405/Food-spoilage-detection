import docx

def update_docx():
    file_path = "report/Banana_Spoilage_Detectionaaaaa.docx"
    doc = docx.Document(file_path)
    
    new_paras = [
        "The spoilage detection system employs a discrete, two-stage decision-making process to evaluate the environmental safety of the storage container in real-time. Rather than waiting for the mathematical Remaining Shelf Life (RSL) to reach zero, the system continuously cross-references raw sensor telemetry against established biological thresholds. This acts as an early-warning \"Traffic Light\" mechanism, generating actionable alerts for warehouse operators before irreversible spoilage occurs.",
        "In the first stage, the rule engine evaluates individual sensor conditions to build a list of contributing factors. These thresholds are defined as follows:",
        "• Methane (CH4) > 0.5 ppm (Critical): Indicates that anaerobic decomposition has initiated. This is an irreversible biological process signifying active rot.",
        "• Temperature > 25°C and <= 30°C (Warning): Elevated ambient temperature is causing the fruit's respiration rate to accelerate, prematurely advancing the biological clock.",
        "• Temperature > 30°C (Critical): The environment is dangerously hot for storage, risking rapid thermal degradation and extreme stress.",
        "• Temperature < 13°C (Warning): Storage temperature has dropped below the safe threshold, introducing a high risk of chilling injury and epidermal peel browning.",
        "• Relative Humidity < 80% (Warning): Suboptimal moisture levels are present, leading to accelerated transpirational moisture loss and physical drying of the fruit.",
        "• Carbon Dioxide (CO2) > 50,000 ppm (Warning): Carbon dioxide accumulation is dangerously high. Immediate ventilation is recommended as the atmosphere is nearing the 70,000 ppm toxicity limit.",
        "• Ethylene (C2H4) > 5.0 ppm (Info): The fruit has reached its climacteric peak, signalling that the autocatalytic ripening phase is accelerating.",
        "In the second stage of the logic, these flagged factors are aggregated by severity (Info, Warning, Critical). The highest severity factor is then combined with the Machine Learning model's predicted Remaining Shelf Life (RSL) to assign the final operational status colour (Green, Yellow, or Red) to the dashboard. This ensures the user receives both an immediate environmental alert and a clear explanation of the underlying cause."
    ]
    
    # We replace paragraphs 227 to 235
    for i in range(9):
        doc.paragraphs[227 + i].text = new_paras[i]
        
    # We have 1 extra paragraph (index 9 in new_paras)
    # Insert it before paragraph 236
    new_p = doc.paragraphs[236].insert_paragraph_before(new_paras[9])
    
    # Optional: We could apply List Bullet style to the bullet points if we want,
    # but since we manually added the bullet '•', normal style is safe and won't break formatting.
    
    doc.save(file_path)
    print("Done updating the document.")

if __name__ == "__main__":
    update_docx()
