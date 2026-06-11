#!/usr/bin/env python3
"""
Extract text+LaTeX from the existing .docx, write a clean Markdown file,
then use pandoc to produce a proper .docx with:
  - LaTeX → Word Equation Objects (OMML)
  - Proper heading styles
  - Page numbers (roman for front matter, arabic for body)
"""

import re
import subprocess
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from copy import deepcopy
from lxml import etree
import shutil

SRC = Path(__file__).parent / "Banana_Spoilage_Detection_Report_p.docx"
MD_OUT = Path(__file__).parent / "report_clean.md"
DOCX_OUT = Path(__file__).parent / "Banana_Spoilage_Detection_Report_Fixed.docx"
IMG_DIR = Path(__file__).parent / "images"

# ── Step 1: Extract images from the original docx ──────────────────────────
def extract_images(doc, img_dir: Path):
    """Extract all images from doc to img_dir, return list of filenames."""
    img_dir.mkdir(exist_ok=True)
    image_files = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            ext = rel.target_ref.split(".")[-1]
            fname = f"image_{len(image_files)+1}.{ext}"
            (img_dir / fname).write_bytes(img_data)
            image_files.append(fname)
    return image_files


# ── Step 2: Full document → Markdown with LaTeX preserved ──────────────────
def docx_to_markdown(doc) -> str:
    """Convert docx paragraphs to markdown, preserving LaTeX math."""
    lines = []
    
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "Normal"
        text = para.text.strip()
        
        if not text:
            lines.append("")
            continue
        
        # Heading styles
        if style == "Heading 1":
            lines.append(f"\n# {text}\n")
        elif style == "Heading 2":
            lines.append(f"\n## {text}\n")
        elif style == "Heading 3":
            lines.append(f"\n### {text}\n")
        elif style == "List Paragraph":
            lines.append(f"- {text}")
        else:
            # Normal paragraph — just output text as-is (LaTeX stays intact)
            lines.append(text)
        
        lines.append("")
    
    return "\n".join(lines)


# ── Step 3: Create a pandoc reference template for page numbering ──────────
def create_reference_docx(ref_path: Path):
    """Create a minimal reference.docx with page number settings."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    from docx.shared import Pt
    font.size = Pt(12)
    
    # Set up sections for page numbering
    # We'll handle this via pandoc's built-in features instead
    
    doc.save(ref_path)


# ── Step 4: Post-process the pandoc output to add page numbers ─────────────
def add_page_numbers(docx_path: Path):
    """Add page numbers to the docx: roman for initial pages, arabic for body."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree
    
    doc = Document(docx_path)
    
    NSMAP = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    
    # Find the section properties
    sections = doc.sections
    
    if len(sections) >= 1:
        # First section: roman numerals (front matter)
        sect = sections[0]
        sectPr = sect._sectPr
        
        # Add page number type — start with roman lowercase
        pgNumType = etree.SubElement(sectPr, f'{{{NSMAP["w"]}}}pgNumType')
        pgNumType.set(f'{{{NSMAP["w"]}}}fmt', 'lowerRoman')
        pgNumType.set(f'{{{NSMAP["w"]}}}start', '1')
        
        # Add footer with page number
        _add_page_number_footer(doc, sect)
    
    # Look for "Chapter 1" or "1   Introduction" heading to insert section break
    # and switch to arabic numbering
    found_chapter1 = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if ('introduction' in text and ('chapter 1' in text or text.startswith('1'))) or text == 'introduction':
            if para.style and 'heading' in para.style.name.lower():
                found_chapter1 = True
                # Insert a section break before this paragraph
                pPr = para._element.get_or_add_pPr()
                sectPr_new = etree.SubElement(pPr, f'{{{NSMAP["w"]}}}sectPr')
                pgNumType_new = etree.SubElement(sectPr_new, f'{{{NSMAP["w"]}}}pgNumType')
                pgNumType_new.set(f'{{{NSMAP["w"]}}}fmt', 'decimal')
                pgNumType_new.set(f'{{{NSMAP["w"]}}}start', '1')
                break
    
    doc.save(docx_path)


def _add_page_number_footer(doc, section):
    """Add a centered page number to the section's footer."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear existing footer
    for p in footer.paragraphs:
        p.clear()
    
    # Add page number field
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = para.add_run()
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # Add PAGE field
    NSMAP_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    fldChar_begin = etree.SubElement(run._element, f'{{{NSMAP_W}}}fldChar')
    fldChar_begin.set(f'{{{NSMAP_W}}}fldCharType', 'begin')
    
    run2 = para.add_run()
    instrText = etree.SubElement(run2._element, f'{{{NSMAP_W}}}instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' PAGE '
    
    run3 = para.add_run()
    fldChar_end = etree.SubElement(run3._element, f'{{{NSMAP_W}}}fldChar')
    fldChar_end.set(f'{{{NSMAP_W}}}fldCharType', 'end')


# ── Step 5: Re-insert images from original doc ────────────────────────────
def reinsert_images(docx_path: Path, img_dir: Path, image_files: list):
    """
    Find figure caption paragraphs in the output docx and insert the 
    corresponding image before each caption.
    """
    from docx import Document
    from docx.shared import Inches
    
    doc = Document(docx_path)
    
    # Map figure numbers to image files
    fig_image_map = {}
    for i, fname in enumerate(image_files):
        fig_image_map[i + 1] = img_dir / fname
    
    # Find paragraphs that start with "Figure X.Y:" 
    figure_paras = []
    for j, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.match(r'Figure\s+(\d+)\.(\d+)', text)
        if match:
            fig_num = int(match.group(2))  # sub-figure number
            chapter = int(match.group(1))
            # Calculate overall figure index
            figure_paras.append((j, para, len(figure_paras) + 1))
    
    # Insert images before each figure caption
    img_idx = 0
    for j, para, fig_overall in figure_paras:
        if img_idx < len(image_files):
            img_path = img_dir / image_files[img_idx]
            if img_path.exists():
                # Insert a new paragraph before the caption paragraph
                new_para = para.insert_paragraph_before("")
                new_para.alignment = 1  # CENTER
                run = new_para.add_run()
                run.add_picture(str(img_path), width=Inches(5.5))
            img_idx += 1
    
    doc.save(docx_path)


# ── Main ──────────────────────────────────────────────────────────────────
def main():
    print("Step 1: Reading original document...")
    doc = Document(SRC)
    
    print("Step 2: Extracting images...")
    image_files = extract_images(doc, IMG_DIR)
    print(f"   Extracted {len(image_files)} images")
    
    print("Step 3: Converting to Markdown with LaTeX preserved...")
    md_content = docx_to_markdown(doc)
    MD_OUT.write_text(md_content, encoding="utf-8")
    print(f"   Written to {MD_OUT}")
    
    print("Step 4: Running pandoc (Markdown+LaTeX → DOCX with OMML equations)...")
    cmd = [
        "pandoc",
        str(MD_OUT),
        "-f", "markdown+tex_math_dollars",
        "-t", "docx",
        "-o", str(DOCX_OUT),
        "--standalone",
        # Use math rendering
        "--mathml",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"   Pandoc error: {result.stderr}")
        return
    print(f"   Generated {DOCX_OUT}")
    
    print("Step 5: Adding page numbers (roman → arabic)...")
    add_page_numbers(DOCX_OUT)
    
    print("Step 6: Re-inserting images...")
    reinsert_images(DOCX_OUT, IMG_DIR, image_files)
    
    print(f"\n✅ Done! Fixed report saved to:\n   {DOCX_OUT}")


if __name__ == "__main__":
    main()
