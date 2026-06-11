import docx

doc = docx.Document("report/Banana_Spoilage_Detectionaaaaa.docx")

i = 226
print(f"\n--- MATCH AT INDEX {i} ---")
for j in range(0, 30):
    if i + j < len(doc.paragraphs):
        p = doc.paragraphs[i+j]
        print(f"[{i+j}] {p.text[:100]}... (Style: {p.style.name})")

